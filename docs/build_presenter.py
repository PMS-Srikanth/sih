# -*- coding: utf-8 -*-
"""
Presenter's Handbook — everything a team member needs to stand up and present
the six-slide deck, plus a comparison against existing solutions and a full
question bank.

    python -m pip install python-docx
    python docs/build_presenter.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C = lambda h: RGBColor.from_string(h)
INK, INK2, INK3 = C("15191C"), C("46525A"), C("737F87")
NAVY, TEAL, GREEN, RED, INDIGO = C("1F4E79"), C("0A6E78"), C("15803D"), C("A32B20"), C("4A42B8")
AMBER = C("B4610F")

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.75)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.12


def shade(cell, hexcolor):
    el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def para(text="", size=10.5, bold=False, italic=False, color=INK2, space=5,
         align=None, mono=False, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if align: p.alignment = align
    for seg in (text if isinstance(text, list) else [text]):
        t, o = seg if isinstance(seg, tuple) else (seg, {})
        r = p.add_run(t)
        r.font.size = Pt(o.get("size", size))
        r.bold = o.get("bold", bold)
        r.italic = o.get("italic", italic)
        r.font.color.rgb = o.get("color", color)
        r.font.name = "Consolas" if o.get("mono", mono) else "Calibri"
    return p


def h1(text, newpage=True):
    if newpage: doc.add_page_break()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(17); r.bold = True; r.font.color.rgb = NAVY
    b = doc.add_paragraph(); b.paragraph_format.space_after = Pt(9)
    br = b.add_run("─" * 66); br.font.size = Pt(7); br.font.color.rgb = C("BBC7D1")


def h2(text, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size = Pt(12.5); r.bold = True; r.font.color.rgb = color


def h3(text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = color


def bullet(text, color=INK2, head=None, indent=0.0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.26 + indent)
    p.paragraph_format.space_after = Pt(2)
    if head:
        r = p.add_run(head); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = INK
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = color


def table(headers, rows, widths=None, size=9, head_fill="1F4E79"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, head_fill)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        r = p.add_run(htxt); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = C("FFFFFF")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            txt, col, bold = (val if isinstance(val, tuple) else (val, INK2, False))
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run(txt); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = col
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def callout(title, body, color=NAVY, fill="EEF3F7"):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade(c, fill)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = color
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body); r2.font.size = Pt(10); r2.font.color.rgb = INK2
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def qa(q, a, tag=None, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q.  " + q); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = color
    if tag:
        rt = p.add_run("   [" + tag + "]"); rt.bold = True; rt.font.size = Pt(8)
        rt.font.color.rgb = INK3
    para(a, indent=0.22, space=3)


CHARTS = r"c:\Users\srika\OneDrive\Desktop\sih-p\docs\charts"


def chart(name, width=6.6, caption=None):
    """Drop in a generated chart. Run docs/build_charts.py first."""
    path = os.path.join(CHARTS, name)
    if not os.path.exists(path):
        para(f"[chart missing: {name} — run  python docs/build_charts.py ]",
             color=RED, italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        para(caption, size=9, color=INK3, align=WD_ALIGN_PARAGRAPH.CENTER, space=9)


def row_explain(claim, plain, matters, push):
    """One comparison row, unpacked so a presenter can improvise around it."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("\u25B8  " + claim)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = TEAL

    for label, text, col in (("In plain words:  ", plain, INK),
                             ("Why it matters:  ", matters, INK),
                             ("If they push back:  ", push, INK)):
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Inches(0.24)
        q.paragraph_format.space_after = Pt(2)
        a = q.add_run(label); a.bold = True; a.font.size = Pt(10); a.font.color.rgb = INK3
        b = q.add_run(text); b.font.size = Pt(10); b.font.color.rgb = col


def slide_head(title, clock, seconds, colour=NAVY):
    """Slide heading with its wall-clock window and budget."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title); r.font.size = Pt(12.5); r.bold = True; r.font.color.rgb = colour
    r2 = p.add_run(f"     {clock}")
    r2.font.size = Pt(11); r2.bold = True; r2.font.color.rgb = AMBER
    r3 = p.add_run(f"   ({seconds})")
    r3.font.size = Pt(9.5); r3.font.color.rgb = INK3


def cover(text):
    """What this slide has to get across. A brief, not a script."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; shade(c, "E9F2EC")
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
    r = p.add_run("COVER HERE:  ")
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GREEN
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


# ══════════════════════════════════════════════════════════════════ COVER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SMART INDIA HACKATHON 2026  ·  SIH26171")
r.font.size = Pt(12); r.bold = True; r.font.color.rgb = RED

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("PRESENTER'S HANDBOOK")
r.font.size = Pt(30); r.bold = True; r.font.color.rgb = NAVY

para("CORDON — on-device visual perception for light-weight browser agents",
     size=13.5, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space=3)
para("Slide-by-slide notes  ·  competitor comparison  ·  question bank",
     size=10.5, color=INK3, align=WD_ALIGN_PARAGRAPH.CENTER, space=14)

callout("READ THIS FIRST",
        "This handbook maps one-to-one onto the six-slide deck. Section 2 is what to SAY for each "
        "slide, with timings. Section 3 is every number you may quote and where it comes from — do "
        "not quote a figure that is not in that table. Sections 5 and 6 are the question bank: "
        "user-facing questions and technical ones. Section 7 lists what we have NOT built, and you "
        "must be willing to say it out loud.", NAVY)

h2("The pitch, in one breath", NAVY)
para([("An AI agent has to see your screen to help you — and today that means uploading it. ",
       {"color": INK}),
      ("Cordon reads your screen on your own device, replaces every personal value with a typed "
       "reference like ", {"color": INK}),
      ("EMAIL_1", {"mono": True, "bold": True, "color": TEAL}),
      (", and lets the server plan the work without ever receiving a single value.", {"color": INK})])

h2("The three sentences to memorise", NAVY)
bullet("A 1.2 MB vision model reads the page inside the browser.", head="1.  ")
bullet("Personal values are swapped for opaque typed handles before anything is sent.", head="2.  ")
bullet("An independent verifier re-reads the outgoing bytes and can refuse to transmit.", head="3.  ")

# ══════════════════════════════════════════════════════════════════ 1
h1("1 · Before you walk in")

h2("Who says what")
table(["Role", "Owns", "Must be able to answer"],
      [[("Presenter", INK, True), "Slides 1, 2 and 5 — the story and the impact",
        "Why this matters; who it helps; why it is not an existing product"],
       [("Technical lead", INK, True), "Slides 3 and 4 — architecture and feasibility",
        "How the model runs in a browser; what the verifier does; the risks"],
       [("Demo operator", INK, True), "The laptop", "Runs the live demo; opens the privacy receipt"],
       [("Everyone", INK, True), "The numbers in Section 3", "Every figure on every slide"]],
      [1.4, 2.7, 3.2])

h2("The one-minute setup, before anyone is watching")
para("Three terminals, one browser. Do this while the previous team presents:")
bullet("Terminal 1 — npm run build, then reload the extension and check the build stamp matches.")
bullet("Terminal 2 — npm run server, left visible. The judge will watch this.")
bullet("Terminal 3 — npm run demo.")
bullet("Chrome — open http://127.0.0.1:8788/application.html and pin the Cordon icon.")
bullet("Side panel — unlock My data. A locked vault means the fill demo does nothing.")

callout("THE ONE THING THAT BREAKS DEMOS",
        "If the build stamp in the side panel header does not match what npm run build printed, the "
        "extension was not reloaded and you are demoing stale code. Check it every single time.", RED,
        "FBEBE9")

# ══════════════════════════════════════════════════════════════════ 2
h1("2 · The nine-minute pitch")

callout("HOW THIS SECTION WORKS",
        "Every slide has a wall-clock window, a time budget, and one sentence in a green box that "
        "MUST come out of your mouth. If you are running late, cut the detail and say the green box. "
        "If you say nothing else on a slide, say that.", NAVY)

h2("Plan A \u2014 with the live demo  (recommended)", NAVY)
para("Use this one. A working prototype is our single biggest advantage over other teams at the "
     "idea stage, and ninety seconds of it beats three minutes of description.")
table(["Clock", "Slide", "Budget", "The job of this segment"],
      [[("0:00 \u2013 0:20", AMBER, True), "1 \u00b7 Title", "20 sec", "State the problem statement. Nothing more."],
       [("0:20 \u2013 2:30", AMBER, True), "2 \u00b7 Idea title", ("2 min 10", INK, True),
        "The problem, then the whole mechanism in four panels."],
       [("2:30 \u2013 4:20", AMBER, True), "3 \u00b7 Technical approach", ("1 min 50", INK, True),
        "How it actually runs. Point at the tech, dwell on the diagram."],
       [("4:20 \u2013 6:00", AMBER, True), ("LIVE DEMO", GREEN, True), ("1 min 40", GREEN, True),
        "Three tasks. Show the receipt. Show the masked face."],
       [("6:00 \u2013 7:15", AMBER, True), "4 \u00b7 Feasibility", "1 min 15",
        "It already runs \u2014 and here is what is not finished."],
       [("7:15 \u2013 8:30", AMBER, True), "5 \u00b7 Impact", "1 min 15",
        "One real person. Then national alignment."],
       [("8:30 \u2013 8:50", AMBER, True), "6 \u00b7 References", "20 sec",
        "Prior art, and the gap we filled."],
       [("8:50 \u2013 9:00", AMBER, True), ("Close", INK, True), "10 sec", "The three-sentence ending."]],
      [1.15, 1.75, 0.85, 3.5], size=9.5)

h2("Plan B \u2014 slides only", NAVY)
para("Use this if you are told slides only, or the laptop dies. Same nine minutes, redistributed.")
table(["Clock", "Slide", "Budget"],
      [[("0:00 \u2013 0:25", AMBER, True), "1 \u00b7 Title", "25 sec"],
       [("0:25 \u2013 3:20", AMBER, True), "2 \u00b7 Idea title", "2 min 55"],
       [("3:20 \u2013 5:40", AMBER, True), "3 \u00b7 Technical approach", "2 min 20"],
       [("5:40 \u2013 7:15", AMBER, True), "4 \u00b7 Feasibility", "1 min 35"],
       [("7:15 \u2013 8:35", AMBER, True), "5 \u00b7 Impact", "1 min 20"],
       [("8:35 \u2013 9:00", AMBER, True), "6 \u00b7 References + close", "25 sec"]],
      [1.4, 3.2, 1.3], size=9.5)

callout("PACE CHECK \u2014 GLANCE AT THE CLOCK AT THESE THREE MOMENTS",
        "At 2:30 you should be leaving slide 2. At 6:00 you should be leaving the demo. At 8:30 you "
        "should be on the last slide. If you are more than twenty seconds behind at any of those, "
        "drop to the green boxes and keep moving \u2014 running over is worse than saying less.", RED,
        "FBEBE9")

# ═══════════════════════════════════════════════ SLIDE 1
slide_head("Slide 1 \u2014 Title page", "0:00 \u2013 0:20", "20 seconds")
cover("The problem statement ID, its title, and your team. Do not touch the solution yet.")
para("Read the ID and the title. Do not editorialise, do not preview the solution. The panel is "
     "still settling; give them a clean twenty seconds to find your slide.", indent=0.15)

# ═══════════════════════════════════════════════ SLIDE 2
slide_head("Slide 2 \u2014 Idea title", "0:20 \u2013 2:30", "2 min 10 \u2014 the most important slide")
cover("Why today’s agents leak, then the four-panel example end to end, then the two innovation columns — and finish on the green box that proves it already runs.")

para("Suggested phrasing below \u2014 your own words will sound better. "
     "What matters is that the brief above gets covered.", size=9.5, color=INK3, italic=True)
h3("0:20 \u2013 0:45  \u00b7  Open with the problem, not the product")
para("\u201cAn AI agent has to see your screen to help you. Every browser agent today solves that by "
     "uploading the screen \u2014 your passwords, your face, everything. That single fact is the "
     "ceiling on what anyone will let an agent do.\u201d", indent=0.2)

h3("0:45 \u2013 1:40  \u00b7  Walk the four panels, left to right, pointing at each")
bullet("\u201cThis is what you see. Your email, your password, your photo.\u201d", head="What you see:  ")
bullet("\u201cThis is everything we send. The email became EMAIL_1. The password was removed \u2014 "
       "not masked, never read. The face was blacked out in the pixels.\u201d", head="What we send:  ")
bullet("\u201cThe server plans the step using EMAIL_1. It never learns the address.\u201d",
       head="Server replies:  ")
bullet("\u201cAnd your browser swaps the handle back and types the real value.\u201d",
       head="Your browser:  ")

h3("1:40 \u2013 1:50  \u00b7  Pause, then land the line under the diagram")
para("\u201cThe server plans the work. It never receives a single value.\u201d", indent=0.2)
para("Actually pause before this sentence. It is the thesis of the entire project and it needs a "
     "beat of silence in front of it.", indent=0.2, italic=True, color=INK3)

h3("1:50 \u2013 2:15  \u00b7  The two columns \u2014 one line each, do not read them all")
para("\u201cThat addresses the problem statement directly: the screen is read on the device, only "
     "structure leaves, and a verifier can refuse to send. What is new is the handle itself \u2014 "
     "because it is stable, the server can tell two fields hold the SAME email without ever learning "
     "it.\u201d", indent=0.2)

h3("2:15 \u2013 2:30  \u00b7  Close on the green box")
para("\u201cAnd this is not a concept. It runs today. A 1.2 megabyte model, 92 percent of every "
     "screen never analysed, perfect precision and recall on our PII test set, and most steps send "
     "zero bytes.\u201d", indent=0.2)

# ═══════════════════════════════════════════════ SLIDE 3
slide_head("Slide 3 \u2014 Technical approach", "2:30 \u2013 4:20", "1 min 50")
cover("The stack in one sweep, then the architecture diagram and where the red boundary sits, then why 92% of a screen is never analysed — and that every claim is reproducible.")

h3("2:30 \u2013 3:00  \u00b7  The stack \u2014 point, do not read")
para("\u201cA TypeScript extension. ONNX Runtime Web on WebGPU with a CPU fallback. A 1.2 megabyte "
     "face detector running inside the browser. And on the server, an open-weights model \u2014 "
     "Llama or Qwen through vLLM \u2014 so it is offline deployable, as the problem statement "
     "requires.\u201d", indent=0.2)

h3("3:00 \u2013 3:50  \u00b7  The architecture diagram \u2014 this is where the time goes")
bullet("\u201cContent script reads the page. Service worker holds the vault \u2014 the web page has "
       "no route to it at all. And the model lives in an offscreen document, because WebGPU is not "
       "available to service workers.\u201d")
bullet("\u201cEverything below this red line is about ten kilobytes of structure and handles. No "
       "values. Ever.\u201d")

h3("3:50 \u2013 4:10  \u00b7  The efficiency argument \u2014 judges score this at 20 percent")
para("\u201cWe do not run the model on every frame. We work out which pixels the page structure "
     "already explains, and look only at what is left \u2014 images, canvases, anything the DOM "
     "cannot describe. That is how 92 percent of a screen is never analysed at all.\u201d", indent=0.2)

h3("4:10 \u2013 4:20  \u00b7  The proof box, bottom right")
para("\u201cEvery claim on this slide is reproducible. One command prints our precision and recall; "
     "another verifies the model loads and its outputs match our code.\u201d", indent=0.2)

# ═══════════════════════════════════════════════ DEMO
slide_head("LIVE DEMO", "4:20 \u2013 6:00", "1 min 40 \u2014 rehearse this until it is muscle memory", GREEN)
cover("Three tasks — one that never leaves the device, one that redacts, and the payload inspector — then point at the face the model masked.")

para("Three tasks only. Do not improvise, do not browse, do not explain the UI. Have the page and "
     "the panel already open before you start talking.", indent=0.15)

table(["Clock", "Do this", "Say this"],
      [[("4:20", AMBER, True), ("Task: save draft", INK, True),
        "\u201cWatch the server terminal. Nothing. That step never left the machine.\u201d"],
       [("4:45", AMBER, True), ("Task: what sensitive data is on this page?", INK, True),
        "\u201cEvery personal value on the page is now a placeholder.\u201d"],
       [("5:10", AMBER, True), ("Expand \u2018What was sent\u2019", INK, True),
        "\u201cThis is the literal text we transmitted. Read it \u2014 the email and the password "
        "are not in there.\u201d"],
       [("5:35", AMBER, True), ("Point at the masked face on the page", INK, True),
        "\u201cAnd the model found that face in the pixels. The page only says it is an image.\u201d"]],
      [0.75, 2.5, 3.95], size=9.5)

callout("IF THE DEMO FAILS \u2014 DECIDE IN FIVE SECONDS, DO NOT FIGHT IT",
        "Say: \u201cLet me show you the same thing without the browser,\u201d switch to the terminal "
        "and run npm run eval. It prints precision, recall, the checksum arbitration and the verifier "
        "results in about two seconds. You lose twenty seconds and no credibility. Wrestling a laptop "
        "for a minute loses both.", RED, "FBEBE9")

# ═══════════════════════════════════════════════ SLIDE 4
slide_head("Slide 4 \u2014 Feasibility and viability", "6:00 \u2013 7:15", "1 min 15")
cover("Why feasibility is easy here (it is already built), then sustainability, then the risks column — and volunteer the Firefox gap out loud.")

h3("6:00 \u2013 6:30  \u00b7  Feasibility is easy for us, so say why")
para("\u201cFeasibility is straightforward here because it is already built. The model is smaller "
     "than most web fonts. It needs no new hardware. And there is no third-party code inside the "
     "privacy boundary at all.\u201d", indent=0.2)

h3("6:30 \u2013 6:50  \u00b7  Sustainability \u2014 a scored criterion most teams skip")
para("\u201cIt is built entirely on open standards, the server model is open-weights so there is no "
     "vendor lock-in, and perception costs us no GPU because it runs on the client.\u201d", indent=0.2)

h3("6:50 \u2013 7:15  \u00b7  The risks column \u2014 do NOT skip this")
para("\u201cWe have named seven risks and what we do about each. And the last one is honest: our "
     "Firefox build compiles from the same codebase, but we have not loaded it yet.\u201d", indent=0.2)
callout("WHY YOU VOLUNTEER THE GAP",
        "Judges are testing whether you know your own weaknesses. A team that names a gap before "
        "being asked is believed on everything else. A team caught hiding one is believed on nothing. "
        "This costs you eight seconds and buys the whole panel.", NAVY)

# ═══════════════════════════════════════════════ SLIDE 5
slide_head("Slide 5 \u2014 Impact and benefits", "7:15 \u2013 8:30", "1 min 15")
cover("One real person on a government portal, widen briefly to the other audiences, then land DPDP and Atmanirbhar Bharat.")

h3("7:15 \u2013 7:50  \u00b7  One real person, not four categories")
para("\u201cTake someone applying for a scholarship on a government portal. Today, letting an agent "
     "fill that form means uploading their Aadhaar. With Cordon the number never leaves the phone. "
     "The server is told only that an Aadhaar-shaped value exists, and which box it belongs in.\u201d",
     indent=0.2)

h3("7:50 \u2013 8:10  \u00b7  Then widen \u2014 quickly")
para("\u201cThe same holds for banking forms, for hospitals and legal offices, and for blind users "
     "who need an agent to read the page for them without surrendering the screen.\u201d", indent=0.2)

h3("8:10 \u2013 8:30  \u00b7  National alignment \u2014 land this one properly")
para("\u201cUnder the DPDP Act, data minimisation is usually a policy promise. Here it is the "
     "transport format \u2014 the server is structurally incapable of receiving the data. And "
     "because the model is open-weights and self-hosted, there is no foreign API anywhere in the "
     "loop.\u201d", indent=0.2)

# ═══════════════════════════════════════════════ SLIDE 6
slide_head("Slide 6 \u2014 Research and references", "8:30 \u2013 8:50", "20 seconds")
cover("Do not read the list. Make one point only: what the prior art does, and the gap none of them fill.")
para("Do not read the list. It exists to prove you did the reading. Use the twenty seconds for one "
     "point only:", indent=0.15)
para("\u201cWe surveyed the prior art. Computer-use agents upload the screen. Redaction tools work "
     "server-side, after the data has already left. Password managers fill locally but cannot reason. "
     "Nobody puts a verifier in front of the network call \u2014 and that gap is what we built.\u201d",
     indent=0.2)

# ═══════════════════════════════════════════════ CLOSE
slide_head("Close", "8:50 \u2013 9:00", "10 seconds", GREEN)
cover("One sentence restating the boundary, then thank the panel.")
para("\u201cEvery browser agent today asks you to upload your screen. We built the one that does not "
     "have to \u2014 so the server does its job on a page it has genuinely never seen. Thank "
     "you.\u201d", indent=0.2)

# ══════════════════════════════════════════════════════════════════ 3
h1("3 · Every number you may quote")

callout("RULE",
        "If a figure is not in this table, do not say it. Every row below is either measured by a "
        "command anyone can run, or a property of a file in the repository. Inventing a number is the "
        "fastest way to lose a panel's trust.", RED, "FBEBE9")

h2("Measured — reproducible on the spot")
table(["Figure", "What it means", "How to prove it"],
      [[("1.21 MB", TEAL, True), "Size of the on-device vision model", "ls the file: ultraface-320.onnx"],
       [("92.0%", TEAL, True), "Share of a screen the model never analyses", "npm run eval"],
       [("1.000 / 1.000", TEAL, True), "PII precision and recall on our test set", "npm run eval"],
       [("5 of 9", TEAL, True), "Benchmark tasks resolved with zero network calls", "npm run eval"],
       [("84.8 ms", TEAL, True), "Model inference, CPU single-thread (faster on WebGPU)",
        "npm run model-check"],
       [("4,420 \u2192 \u226424", TEAL, True), "Candidate boxes reduced by NMS", "npm run model-check"],
       [("310,000", TEAL, True), "PBKDF2 iterations protecting the local vault", "crypto.ts"],
       [("~10 KB", TEAL, True), "Typical sanitized payload per step", "Side panel, bytes sent"],
       [("0 bytes", TEAL, True), "Sent when a step resolves on the device", "Side panel + server log"]],
      [1.25, 3.15, 2.85])

h2("Architectural — true by construction, not benchmarked")
table(["Claim", "Why it is true"],
      [[("Passwords are never read into the payload", INK, True),
        "The redactor emits sensitive:true and never touches the value. Not masking — absence."],
       [("Masks are irreversible", INK, True),
        "Black is composited INTO the bitmap and re-encoded. Original pixels stop existing."],
       [("The page cannot reach the vault", INK, True),
        "Zero storage references in content scripts; the vault lives in the service worker."],
       [("Exactly one network call exists", INK, True),
        "One fetch in the extension, and the verifier stands in front of it."],
       [("No third-party runtime code in the boundary", INK, True),
        "ONNX Runtime Web is the only dependency, and it sits in the offscreen document."]],
      [2.5, 4.75])

# ══════════════════════════════════════════════════════════════════ 4
h1("4 · How Cordon compares to what exists")

callout("HOW TO USE THIS SECTION",
        "Each comparison is a chart followed by a plain-language explainer. Learn the EXPLAINERS, "
        "not the charts — if you understand why each row is true, you can answer a question that is "
        "worded differently. Reciting a table cannot survive a follow-up question.", NAVY)

callout("AND BE HONEST ABOUT WHAT THIS IS",
        "These are ARCHITECTURAL comparisons: what each design can and cannot do, taken from public "
        "documentation. We have NOT run anybody else's product and we have not benchmarked anyone. "
        "If a judge presses, say exactly that: \u201cThis compares what each approach is capable of "
        "by construction, not lab results.\u201d Never invent a competitor's number.", RED, "FBEBE9")

h2("Who we are comparing against", NAVY)
table(["The four approaches", "What it is", "Real examples"],
      [[("CORDON \u2014 ours", TEAL, True), "Reads the screen on your device, sends only structure",
        "This project"],
       [("Computer-use agents", INK, True), "Take a screenshot, send it to a model, act on the reply",
        "Browser-automation agents, RPA tools"],
       [("Server PII tools", INK, True), "Detect and mask personal data \u2014 on a server",
        "Presidio-class redaction services"],
       [("Password managers", INK, True), "Store credentials locally and autofill them",
        "Browser and third-party password tools"]],
      [1.9, 3.1, 2.2])

# ═══════════════ chart 1
h2("A · How much of you leaves the machine", NAVY)
chart("01_payload.png", 6.5,
      "Typical payload per step. Log scale \u2014 each gridline is ten times the previous one.")

row_explain(
    "We send about 10 KB. A screenshot agent sends 200 KB to 2 MB.",
    "We send a list of what is on screen \u2014 roles, labels, positions \u2014 plus placeholders. "
    "They send the actual picture of your screen.",
    "It is not just bandwidth. A screenshot contains everything visible: your name, the email in "
    "your inbox behind the form, the document open in the next tab. Our list contains none of it "
    "because it was never a picture in the first place.",
    "\u201cCouldn't they just compress the screenshot?\u201d \u2014 Compression changes the size, not "
    "what is in it. A 50 KB screenshot still shows your face.")

# ═══════════════ chart 2
h2("B · What actually crosses the boundary", NAVY)
chart("02_what_leaves.png", 6.9,
      "Green means it never leaves your machine. Red means it does.")

row_explain(
    "Your typed values never leave. Theirs do.",
    "When you type your email into a form, we replace it with the label EMAIL_1 before anything is "
    "sent. A screenshot agent sends the pixels showing your email.",
    "This is the whole idea in one row. The server can still plan \u2014 it knows there is an email "
    "field and which one \u2014 but it cannot read, store, leak or be subpoenaed for a value it "
    "never received.",
    "\u201cSo the server is working blind?\u201d \u2014 Not blind, anonymous. It sees the shape of "
    "the page perfectly. It just cannot see the contents of the boxes.")

row_explain(
    "Your password is never even read.",
    "Most systems mask a password \u2014 they read it, then hide it. We never read it. The payload "
    "says the field is sensitive and stops there.",
    "You cannot leak what you never held. There is no masked copy sitting in memory, no debug log "
    "with it in, no bug that could accidentally include it.",
    "\u201cHow does it log me in then?\u201d \u2014 It does not. Password autofill is off by "
    "default; you type it, or your password manager does. The agent handles everything around it.")

row_explain(
    "Faces and ID documents are blacked out in the pixels themselves.",
    "The model finds the face on your machine and paints solid black over that rectangle, into the "
    "image, before the image is re-encoded.",
    "The original pixels stop existing in the file we send. This is different from drawing a box on "
    "top, or blurring \u2014 both of those can be undone.",
    "\u201cWhy not blur, like the problem statement suggests?\u201d \u2014 Blur and pixelation can "
    "often be partially reversed with image processing. Solid black destroys the information. It is "
    "a deliberate upgrade on the example, not a shortcut.")

# ═══════════════ chart 3
h2("C · WHEN the data gets cleaned \u2014 the part people miss", NAVY)
chart("03_timeline.png", 6.7,
      "The dashed line is the moment data leaves your computer. Position on the line is what matters.")

row_explain(
    "We redact BEFORE the network call. Server tools redact after.",
    "A server-side redaction tool has to receive your data in order to clean it. By the time it "
    "acts, the data has already crossed the internet and landed on someone else's machine.",
    "Their protection is about what gets STORED. Ours is about what gets SENT. If the connection is "
    "intercepted, or that server is breached, or the operator is compelled to hand over logs, "
    "server-side redaction protects nothing.",
    "\u201cSurely the connection is encrypted anyway?\u201d \u2014 TLS protects it in transit. It "
    "does not protect it once it arrives, and it does nothing about who runs the server.")

# ═══════════════ chart 4
h2("D · Which part of the job runs where", NAVY)
chart("06_workload.png", 6.5,
      "Only planning ever reaches the server \u2014 and only when the device cannot decide alone.")

row_explain(
    "Reading, detecting and redacting are 100% on your device.",
    "The screen is read locally. Personal data is found locally. Redaction happens locally. Only "
    "the decision about what to do next may go out.",
    "This is what the problem statement asks for in so many words: the local model reads the screen "
    "and takes the decision, and sends to the server only if it must.",
    "\u201cWhy send anything at all?\u201d \u2014 Because the statement itself says a laptop cannot "
    "host a full reasoning pipeline. We put exactly the part that needs weight on the server, and "
    "nothing else.")

row_explain(
    "In our benchmark, 5 of 9 tasks needed no server at all.",
    "\u201cClick Submit\u201d is unambiguous. The device just does it \u2014 no network, no "
    "redaction work, about fifteen milliseconds.",
    "It is a privacy win and a speed win at once. The steps that never leave cannot leak, and they "
    "cost no round trip.",
    "\u201cIs 5 of 9 representative?\u201d \u2014 It is our benchmark set of nine tasks, and we say "
    "so. The proportion moves with the task; the mechanism does not.")

# ═══════════════ chart 5
h2("E · Safety features, present or absent", NAVY)
chart("05_safety.png", 6.6,
      "Six safety behaviours. Most alternatives have none of them.")

row_explain(
    "A verifier that can refuse to send.",
    "A separate piece of code, sharing nothing with the redactor, re-reads the exact bytes about to "
    "go out and can stop the send.",
    "Redaction is code and code has bugs. Everyone else logs what they redacted. We check it "
    "afterwards, independently, and abort if it is wrong.",
    "\u201cWhat if the verifier has a bug too?\u201d \u2014 Then it fails closed: on any failure it "
    "re-redacts harder, re-checks twice, and if still unsure refuses to transmit at all.")

row_explain(
    "A receipt for every step, showing the exact bytes.",
    "Open the side panel and read the literal JSON we sent. Not a summary \u2014 the actual text.",
    "It converts a privacy claim into something checkable. A judge does not have to believe us; "
    "they can read it.",
    "\u201cCouldn't you show a fake receipt?\u201d \u2014 The same payload goes to the server, and "
    "the server terminal logs its size and content class. Compare the two on screen.")

row_explain(
    "Prompt-injection defence that does not rely on the model behaving.",
    "A hostile page can hide text telling the AI to leak your email. Our defence is that each value "
    "type is bound to the kinds of field it may enter \u2014 an email cannot go into a comment box.",
    "Everyone else defends by wording the prompt more firmly. That is asking the model not to be "
    "fooled. Ours works even when it is fooled, because the client refuses.",
    "\u201cWhat if the page mislabels a comment box as an email field?\u201d \u2014 A large "
    "free-text area is disqualified as a PII sink regardless of its label.")

row_explain(
    "It confirms the value actually landed.",
    "After typing, we read the field back and compare. Websites revert values, truncate them, or "
    "reformat them silently.",
    "An agent that assumes success fills half a form wrongly and submits it. Ours reports the step "
    "as failed instead.",
    "\u201cDoes that mean you keep a copy of the value?\u201d \u2014 Only the verdict travels: "
    "verified yes or no, a reason, and a character count. That is why the panel shows 9/13 chars "
    "rather than the text.")

# ═══════════════ chart 6
h2("F · Our own measured numbers", NAVY)
chart("04_measured.png", 6.7,
      "MEASURED \u2014 reproducible in about two seconds with  npm run eval.")

para([("These four are ours and ours alone \u2014 no competitor is implied. ",
       {"bold": True, "color": INK}),
      ("They are also the only figures on the deck that came from running something rather than "
       "reading documentation. If a judge asks you to prove one, run the command.", {})])

h2("The one-sentence version of this whole section", NAVY)
callout("SAY THIS IF YOU ONLY GET ONE SENTENCE",
        "Everyone else either uploads your screen and hopes the server behaves, or cleans the data "
        "after it has already arrived. We clean it before it leaves, and we check our own work before "
        "we send \u2014 so the server does its job on a page it has genuinely never seen.", TEAL,
        "E9F2EC")

# ══════════════════════════════════════════════════════════════════ 5
h1("5 · Question bank — user-oriented")

para("Non-technical questions, from judges, users, or anyone at the stall. Answer in plain "
     "language. Never say \u201cit's encrypted\u201d and stop there — say what that protects against.")

qa("Is my data actually safe if I type it into a browser extension?",
   "Yes, and here is precisely why. Your details are encrypted on your own machine with AES-256, "
   "under a key made from a passphrase only you know. The web page you are visiting cannot reach it "
   "— that is enforced by the browser, not by us being careful. And it is never included in anything "
   "sent to our server. You can open the panel and read the exact bytes we transmit.", "trust")

qa("Where exactly is my data stored?",
   "In the extension's own encrypted storage on your device. Not in a cloud account, not on our "
   "server, not in a file any website can open. The encryption key lives only in memory and "
   "disappears when you close Chrome, so even the encrypted copy cannot be read until you unlock it "
   "again.", "storage")

qa("What happens if I forget my passphrase?",
   "The data is gone. There is no recovery and no back door — if we could recover it, so could "
   "someone else. The interface warns you about this before you commit to a passphrase.", "storage")

qa("Can I see what the agent typed into the form?",
   "Yes. The side panel has a section called \u201cWhat the agent entered\u201d listing every field it "
   "filled. Values are hidden by default, because a side panel is often on screen while other people "
   "are around, and you press Show to reveal one.", "transparency")

qa("How do I know it really removed my password before sending?",
   "Open the privacy receipt for that step and expand \u201cWhat was sent\u201d. It shows the literal "
   "text that crossed the boundary. Read it — your password is not in there. It is not hidden or "
   "starred out; the field simply says sensitive: true, because the value was never read.", "trust")

qa("What if it fills in the wrong thing, or fills the wrong box?",
   "Two protections. Before typing, the agent checks that the field is the right kind for that value "
   "— an email can only go into something that is actually an email field. After typing, it reads "
   "the field back and compares. If the value did not land correctly the step is marked failed, not "
   "quietly assumed to have worked.", "safety")

qa("Will it submit something without asking me?",
   "No. Anything irreversible — submit, pay, delete, confirm — always stops and asks for your "
   "explicit approval, however confident the model is.", "safety")

qa("Does it work on any website, or only your demo pages?",
   "Any website. The extension runs on all URLs. Our demo pages exist so we can show specific "
   "situations on cue — an ID photo, a blank form — not because it needs them.", "scope")

qa("Do I need internet for this to work?",
   "For part of it, no. Reading the screen, finding personal data and redacting it all happen on your "
   "machine. In our benchmark five of nine tasks finish without any network call at all. Complex "
   "planning does need the server.", "offline")

qa("Is this slower than a normal AI agent?",
   "For simple actions it is much faster, because there is no round trip at all — a click resolves "
   "in about fifteen milliseconds. For complex steps it is comparable, because we send about ten "
   "kilobytes of structure instead of a full screenshot.", "performance")

qa("Who can see my face if I upload a photo?",
   "Nobody outside your device. The model finds the face on your machine and paints solid black over "
   "it, into the actual image, before anything is sent. We do not blur — blur can often be undone.",
   "vision")

qa("What does your company get out of my data?",
   "Nothing, because we never receive it. There is no per-user data stored on the server at all. "
   "That is not a policy we are promising to follow — the system is built so the server cannot "
   "receive it in the first place.", "trust")

qa("Why should I trust you more than a big company's assistant?",
   "Do not trust us — check. Every step produces a receipt with the exact bytes we sent. You can read "
   "it yourself. That is a different kind of assurance from a privacy policy you have to take on "
   "faith.", "trust")

# ══════════════════════════════════════════════════════════════════ 6
h1("6 · Question bank — technical")

para("Expect these from the technical judge. Where a number is involved, quote from Section 3 "
     "only. Where we have not built something, say so.")

qa("Why not run the LLM locally as well? Then nothing leaves at all.",
   "Because the problem statement says the local system cannot host a full pipeline — that is the "
   "premise we are given. Practically, a three-billion-parameter model is roughly two gigabytes of "
   "weights in a browser tab, which would destroy the client-resource metric. So the split is "
   "deliberate: vision on the client, reasoning on the server.", "architecture")

qa("Is it really a Vision Transformer?",
   "No, and we are careful to say so. UltraFace is a convolutional network. The problem statement "
   "asks for a Vision Transformer OR an equivalent computer vision model, and a CNN qualifies. We "
   "chose it because it is 1.2 megabytes rather than a hundred, and metric four grades client "
   "resource use. A ViT-Tiny crop classifier is planned work, not a claim.", "model")

qa("If you already have a vision model, why do you need the DOM at all?",
   "They answer different questions at wildly different costs. An input of type password is certain "
   "and free to identify; recognising a password field from pixels is neither. The DOM handles what "
   "it can explain, and vision runs only on what it cannot — images, canvases, custom widgets. That "
   "is how 92% of a screen goes unanalysed.", "architecture")

qa("How does the server plan anything if it cannot see values?",
   "It sees structure and types. It knows element el_7 is a textbox labelled Email that holds "
   "EMAIL_1, and that el_9 is a password field whose value was removed. Planning \u201cfill the email "
   "box\u201d needs the type, not the value. Because handles are stable, it can also tell that two "
   "fields hold the same email — a relationship it does need.", "protocol")

qa("What stops the model from just asking for the raw email?",
   "It has no channel to. The response schema only accepts a handle or a string the server itself "
   "composed, and the server screens its own output for anything that pattern-matches as real PII. "
   "Even if it asked, the client would have nothing to give — the server never received it.", "security")

qa("How do you defend against prompt injection from a hostile page?",
   "Structurally, not by persuasion. Each handle carries an allowedSinks rule binding its class to "
   "the kinds of field it may be written into. An email handle cannot be typed into a comment box, "
   "and a large free-text area never qualifies as a PII sink. Even a completely fooled model cannot "
   "cause exfiltration, because the client refuses to resolve the handle there.", "security")

qa("How is redaction precision measured, and why does it matter?",
   "Metric three grades precision of redaction, so blanking a paragraph to hide one email costs "
   "marks. We use Range.getClientRects to mask exactly the matched characters, and the detector's own "
   "box plus a four-pixel margin for regions. Our harness reports mask IoU, over-redaction rate and "
   "leak rate.", "metrics")

qa("How do you keep PII precision at 1.000 without hurting recall?",
   "Checksums, mostly. Order number 1234567890 and Phone 9876543210 are identical to a regex. But a "
   "twelve-digit number failing Verhoeff is not an Aadhaar, and a sixteen-digit number failing Luhn "
   "is not a card. That is hard negative evidence at almost no cost, so we can search aggressively "
   "and still prove when we are wrong.", "detection")

qa("What is your fusion rule when detectors disagree?",
   "Noisy-OR over calibrated probabilities, with two thresholds. Above 0.80 we redact; below 0.35 we "
   "keep; between them a context tie-break decides using label proximity, checksum outcome, container "
   "role, repetition and task relevance.", "detection")

qa("Why is the model in an offscreen document rather than the service worker?",
   "Because MV3 service workers do not expose WebGPU, or a DOM, or canvas. The offscreen document is "
   "the only extension context that has them. It receives frames and returns coordinates — it has no "
   "access to the vault and no network access.", "architecture")

qa("What happens when WebGPU is unavailable?",
   "ONNX Runtime Web falls back to WASM with SIMD, and we request the providers in that order. The "
   "fallback is tested, not aspirational. Fast mode completes a task with no vision at all.",
   "runtime")

qa("How does the agent avoid clicking the wrong element after a re-render?",
   "Grounding. Before every action we re-derive a stability signature from the element's role, "
   "accessible name and rounded geometry, and compare it to what the plan was made against. On a "
   "mismatch the action is discarded and the page re-read, rather than clicking whatever now occupies "
   "that position.", "execution")

qa("How do you know a value actually entered the field?",
   "We read it back and compare. Six failure modes are covered and tested: exact match, reformatting "
   "by an input mask, reversion by a framework, truncation by maxlength, a different value entirely, "
   "and a dashed card mask. A fill that did not stick is reported as failed. Only the verdict travels "
   "— never the value.", "execution")

qa("Walk me through the verifier.",
   "Six checks, sharing no code with the redactor. V1 re-scans the serialised bytes with the full "
   "pattern battery. V2 asserts no vault plaintext appears anywhere, including URLs and alt text. V3 "
   "re-samples the masked bitmap to confirm every mask actually painted. V4 sweeps for high-entropy "
   "strings that are not known handles. V5 is a deny-by-default key whitelist. V6 escalates: "
   "re-redact harder, re-verify twice, then abort rather than send.", "security")

qa("What crypto, and what are its limits?",
   "AES-256-GCM under a PBKDF2-SHA256 key at 310,000 iterations. GCM authenticates, so a wrong "
   "passphrase fails cleanly rather than returning plausible garbage. The key sits in session storage "
   "— memory-backed, wiped when the browser closes. The honest limit: an attacker with code execution "
   "inside the extension while it is unlocked could read the key. Against the page, the network and "
   "disk theft, it holds.", "security")

qa("Which server model, and is it really offline deployable?",
   "We speak the OpenAI-compatible chat API, which vLLM, Ollama and llama.cpp all expose. So the same "
   "code runs against Qwen2.5-VL or Llama-3.2-Vision self-hosted, with no vendor SDK. If the model is "
   "unreachable or replies badly, it falls back to a rule-based planner rather than stalling the "
   "agent.", "server")

qa("What is your end-to-end latency?",
   "It depends on the route, and that is the point. A locally resolved step is about fifteen "
   "milliseconds with no network at all. A server step is dominated by the round trip, because we "
   "send roughly ten kilobytes rather than a screenshot. The panel shows per-stage timings live.",
   "performance")

qa("How would this scale to a million users?",
   "Perception, detection and redaction all run on the client, so the fleet only ever sees the "
   "escalated fraction, at about ten kilobytes per step. There is no per-user data stored server-side "
   "— which is also why there is nothing to breach.", "scale")

# ══════════════════════════════════════════════════════════════════ 7
h1("7 · What we have NOT built")

callout("SAY THESE BEFORE YOU ARE ASKED",
        "Volunteering a gap costs you nothing and buys credibility for everything else. Being caught "
        "hiding one costs you the panel.", RED, "FBEBE9")

table(["Not built", "Status", "How to say it"],
      [[("Local NER model", INK, True), ("Phase 3", RED, True),
        "\u201cTwo of our four detectors are live. Names in free prose need NER — that is next.\u201d"],
       [("OCR on image text", INK, True), ("Phase 3", RED, True),
        "\u201cWe detect that a document is present; reading text out of it is the next detector.\u201d"],
       [("ViT-Tiny crop classifier", INK, True), ("Phase 3", RED, True),
        "\u201cWe run a CNN detector today. A ViT classifier on crops is planned.\u201d"],
       [("Firefox build tested", INK, True), ("Untested", RED, True),
        "\u201cIt compiles from the same codebase with its own manifest. We have not loaded it yet.\u201d"],
       [("Labelled set for metric 1", INK, True), ("Open", RED, True),
        "\u201cOur PII numbers are measured. Screen-understanding accuracy needs a labelled corpus — "
        "that is our largest remaining gap.\u201d"],
       [("Full form automation", INK, True), ("Partial", RED, True),
        "\u201cStandard fields fill reliably. Unusual labels with no autocomplete attribute still miss, "
        "and we are extending the classifier.\u201d"]],
      [1.7, 0.9, 4.65])

h2("How to answer \u201cso it does not really work yet?\u201d")
para("\u201cIt works end to end today — you can watch it read a page, find personal data, mask a face, "
     "send a censored map, and act on the reply. What is unfinished is breadth: more detectors, more "
     "browsers, a larger test corpus. We chose to build one complete path rather than four partial "
     "ones, because a working loop can be extended and a scaffold cannot.\u201d", indent=0.15)

# ══════════════════════════════════════════════════════════════════ 8
h1("8 · The live demo")

table(["#", "Do", "Say", "Proves"],
      [["1", ("Open application.html, click Cordon", INK, True),
        "\u201cIt is reading the page now.\u201d", "On-device perception"],
       ["2", ("Task: save draft", INK, True),
        "\u201cWatch the server terminal. Nothing. That step never left the machine.\u201d",
        "Local-first routing"],
       ["3", ("Task: what sensitive data is on this page?", INK, True),
        "\u201cEvery personal value is now a handle.\u201d", "Detection and redaction"],
       ["4", ("Expand \u201cWhat was sent\u201d", INK, True),
        "\u201cRead it. The email and password are not there.\u201d", "The privacy claim, checkable"],
       ["5", ("Replace the photo with a real one", INK, True),
        "\u201cThe DOM only knows there is an image. The model finds the face.\u201d", "The vision model"],
       ["6", ("Unlock My data, open job-form.html", INK, True),
        "\u201cDashed boxes mean my device has a value ready.\u201d", "Local profile vault"],
       ["7", ("Task: fill this form from my profile", INK, True),
        "\u201cThe server decided which slot goes where without seeing any of them.\u201d",
        "Handles end to end"],
       ["8", ("Open \u201cWhat the agent entered\u201d", INK, True),
        "\u201cAnd I can audit exactly what it typed.\u201d", "User transparency"],
       ["9", ("Task: submit application", INK, True),
        "\u201cIrreversible actions always stop for a human.\u201d", "Safety policy"]],
      [0.32, 2.35, 3.05, 1.55], size=9)

callout("IF THE DEMO BREAKS",
        "Do not fight it. Say: \u201cLet me show you the same thing without a browser\u201d and run "
        "npm run eval. It prints the precision and recall, the checksum arbitration, the coverage "
        "map result and the verifier checks in about two seconds. A team with a fallback looks "
        "prepared; a team wrestling a laptop does not.", NAVY)

# ══════════════════════════════════════════════════════════════════ 9
h1("9 · One-page cheat sheet")

h2("If you remember nothing else")
bullet("The server plans the work. It never receives a single value.", head="The line:  ")
bullet("1.2 MB model  ·  92% of screen skipped  ·  1.000 precision and recall  ·  0 bytes most steps",
       head="The numbers:  ")
bullet("Handles keep meaning  ·  vision only where the page is opaque  ·  a verifier that can refuse",
       head="The three firsts:  ")
bullet("Local NER, OCR, Firefox testing, and a labelled screen corpus.", head="The gaps:  ")

h2("Words to use, and words to avoid")
table(["Say this", "Not this", "Why"],
      [[("equivalent CV model", GREEN, True), ("Vision Transformer", RED, True),
        "UltraFace is a CNN. The PS allows equivalents; claiming a ViT invites a question you lose."],
       [("solid masking", GREEN, True), ("blurring", RED, True),
        "Blur is partially reversible. We deliberately do better than the PS example."],
       [("the value was removed", GREEN, True), ("the password is masked", RED, True),
        "It was never read into the payload at all. Absence, not concealment."],
       [("unidentifiable", GREEN, True), ("invisible", RED, True),
        "The server does learn that a field exists. The PS's own word is unidentifiable."],
       [("architectural comparison", GREEN, True), ("benchmark", RED, True),
        "We have not run competitors' code. Say what the comparison actually is."]],
      [1.7, 1.5, 4.05])

h2("Three sentences to close on")
para("\u201cEvery browser agent today asks you to upload your screen. We built the one that does not "
     "have to. The model reads your page on your own machine, your data is replaced with references "
     "before anything is sent, and a verifier checks the bytes and can refuse — so the server plans "
     "your work without ever receiving a single value of yours.\u201d", indent=0.15)

BASE = r"c:\Users\srika\OneDrive\Desktop\sih-p\docs\Cordon_Presenter_Handbook"
target = f"{BASE}.docx"
for i in range(1, 30):
    try:
        doc.save(target); break
    except PermissionError:
        target = f"{BASE}_{i}.docx"
else:
    raise SystemExit("close the document in Word and re-run")
print("saved", os.path.basename(target))
